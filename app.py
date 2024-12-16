import os
import cv2
import pytesseract
import numpy as np
import tensorflow as tf
from flask import Flask, request, render_template, send_from_directory, redirect, url_for
from docx import Document
from docx.shared import Pt
from pdf2image import convert_from_path
from werkzeug.utils import secure_filename
import io
from PIL import Image as PILImage

app = Flask(__name__, static_folder='statics')

# Configuration for directories
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
OUTPUT_FOLDER = os.path.join(BASE_DIR, 'document_output')

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

# Ensure directories exist
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
if not os.path.exists(OUTPUT_FOLDER):
    os.makedirs(OUTPUT_FOLDER)

# Load font classification model and set up class labels
model = tf.keras.models.load_model('models/Fontclassifier1.h5')
class_labels = {
    0: "Khmer OS",
    1: "Khmer OS Battambong",
    2: "Khmer OS Siemreap",
    # Add more classes as needed
}

def preprocess_image(image, target_size=(256, 256)):
    """Preprocess the image for font classification."""
    image = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
    image = cv2.resize(image, target_size)
    image = image / 255.0  # Normalize to [0, 1]
    return image

def detect_font(image):
    """Predict the font of a given image region."""
    processed_image = preprocess_image(image)
    processed_image = np.expand_dims(processed_image, axis=0)  # Add batch dimension
    predictions = model.predict(processed_image)
    predicted_class = np.argmax(predictions, axis=1)[0]
    confidence = np.max(predictions)
    return class_labels[predicted_class], confidence

from docx.shared import Pt, RGBColor

def set_font(run, font_name, font_size=12):
    """Set the font style for a run of text."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)  # Set the font color to black (default)

def process_image(image, doc):
    """Process a single image file, perform OCR, detect fonts, and append results to a .docx file."""
    gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    text_data = pytesseract.image_to_data(gray_image, config=r'-l khm+eng --psm 6', output_type=pytesseract.Output.DICT)

    # Perform OCR and detect font for the entire image
    result_text = ""
    for i in range(len(text_data['text'])):
        if int(text_data['conf'][i]) > 0:  # Filter by confidence
            word = text_data['text'][i].strip()
            if word:
                result_text += word + " "  # Concatenate all words into one string for the full OCR result

    # Detect font for the whole document (based on the first block of text)
    first_word_image = image[0:100, 0:100]  # Take the first part of the image to detect font
    font_name, font_confidence = detect_font(first_word_image)

    # Only consider font detection if confidence is at least 0.4
    if font_confidence < 0.4:
        font_name = "Font not detected"
        font_confidence = "N/A"

    # Save the OCR text to the document with the specified font
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(result_text)
    set_font(run, "Khmer OS Battambang", 12)  # Apply Khmer OS Battambang font

    return result_text, font_name, font_confidence

def process_pdf(pdf_file, output_file_path):
    """Process a PDF file and save OCR + font results to a combined .docx file."""
    doc = Document()
    font_name = None
    font_confidence = None
    
    images = convert_from_path(pdf_file)  # Convert PDF to a list of PIL Image objects
    for image in images:
        image_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)  # Convert PIL to OpenCV image
        result_text, page_font_name, page_font_confidence = process_image(image_cv, doc)  # Pass the image data directly
        
        # If this is the first page, save the font name and confidence
        if font_name is None and font_confidence is None:
            font_name = page_font_name
            font_confidence = page_font_confidence

    doc.save(output_file_path)
    return font_name, font_confidence  # Return font info

# Function to extract text from the docx
def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    doc_text = ""

    for para in doc.paragraphs:
        for run in para.runs:
            doc_text += run.text + " "  # Concatenate the runs to get the full paragraph text

    return doc_text


@app.route('/')
def index():
    error_message = request.args.get('error')  # Get error message if any
    return render_template('index.html', error_message=error_message)

@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return redirect(url_for('index', error="No file uploaded."))

    file = request.files['file']
    if file.filename == '':
        return redirect(url_for('index', error="No file selected."))

    # Save uploaded file
    filename = secure_filename(file.filename)
    upload_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(upload_path)

    # Determine output file path
    output_filename = f"{os.path.splitext(filename)[0]}.docx"
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

    # Process file based on type
    file_ext = os.path.splitext(filename)[1].lower()

    if file_ext == '.pdf':
        # For PDFs, convert pages to images and process
        font_name, font_confidence = process_pdf(upload_path, output_path)
    elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
        # For image files, load and process the image
        try:
            img = PILImage.open(upload_path)  # Open the image using PIL
            image_cv = np.array(img)  # Convert PIL image to NumPy array (OpenCV format)
            if img.mode == 'RGB':  # Ensure the image is in RGB mode
                image_cv = cv2.cvtColor(image_cv, cv2.COLOR_RGB2BGR)
        except Exception as e:
            return redirect(url_for('index', error=f"Error loading image: {str(e)}"))
        
        doc = Document()
        result_text, font_name, font_confidence = process_image(image_cv, doc)  # Pass the OpenCV image here
        doc.save(output_path)
    else:
        return redirect(url_for('index', error="Unsupported file type. Please upload a PDF or an image file."))

    # Extract full text from the .docx file
    doc_text = extract_text_from_docx(output_path)

    # Pass the full text, font name, font confidence, and download link to the template
    return render_template('index.html', doc_text=doc_text, font_name=font_name, font_confidence=font_confidence, download_link=output_filename)

@app.route('/download/<filename>')
def download(filename):
    return send_from_directory(app.config['OUTPUT_FOLDER'], filename, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)